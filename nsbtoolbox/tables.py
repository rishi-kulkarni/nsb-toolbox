import re
from enum import Enum
from functools import lru_cache
from typing import Optional

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches, Pt
from docx.table import _Cell, _Row
from docx.text.paragraph import Paragraph

from .classes import QuestionType, Subject, TossUpBonus, ErrorLogger
from .docx_utils import (
    clear_cell,
    copy_run_formatting,
    fuse_consecutive_runs,
    highlight_cell_text,
    highlight_paragraph_text,
    preprocess_cell,
    shade_columns,
    split_run_at,
)

COL_WIDTHS = (
    0.72,
    0.61,
    4.71,
    0.58,
    0.68,
    0.68,
    0.68,
    0.68,
    0.81,
    0.64,
    4.71,
    1.51,
    1.44,
)


class QuestionFormatterState(Enum):
    Q_START = 0
    STEM_END = 1
    CHOICES = 2
    ANSWER = 3


def make_jans_shadings(table):
    """Shades the 4, 6, 7, and 8th cells in a row the appropriate colors.

    Parameters
    ----------
    row : Document.rows object
    """
    shade_columns(table.columns[3], "#FFCC99")
    shade_columns(table.columns[5], "#e5dfec")
    shade_columns(table.columns[6], "#daeef3")
    shade_columns(table.columns[7], "#daeef3")


def initialize_table(nrows=0, path: Optional[str] = None) -> Document:
    """Initializes a docx file containing the Science Bowl header row.

    Parameters
    ----------
    nrows : int, optional
        Number of extra rows to append to the table, by default 0

    path : Optional[str], optional
        Path that the docx file should be saved to.
        If None, merely returns the document, by default None

    Returns
    -------
    Document
    """
    document = Document()
    font = document.styles["Normal"].font
    font.name = "Times New Roman"
    font.size = Pt(11)

    table = document.add_table(rows=1 + nrows, cols=13)

    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False

    table.cell(0, 0).paragraphs[0].add_run("TUB")
    table.cell(0, 1).paragraphs[0].add_run("Subj")
    ques_run = table.cell(0, 2).paragraphs[0].add_run("Ques")
    ques_run.italic = True
    table.cell(0, 3).paragraphs[0].add_run("LOD")
    table.cell(0, 4).paragraphs[0].add_run("LOD-A")
    table.cell(0, 5).paragraphs[0].add_run("Set")
    table.cell(0, 6).paragraphs[0].add_run("Rd")
    table.cell(0, 7).paragraphs[0].add_run("Q Letter")
    table.cell(0, 8).paragraphs[0].add_run("Author")
    table.cell(0, 9).paragraphs[0].add_run("ID")
    table.cell(0, 10).paragraphs[0].add_run("Source")
    table.cell(0, 11).paragraphs[0].add_run("Comments")
    table.cell(0, 12).paragraphs[0].add_run("Subcat")

    make_jans_shadings(table)

    for idx, column in enumerate(table.columns):
        for cell in column.cells:
            cell.width = Inches(COL_WIDTHS[idx])

    if path is not None:
        document.save(path)

    return document


@lru_cache
def _compile(regex: str):
    return re.compile(regex, re.IGNORECASE)


class QuestionCellFormatter:
    """Formats a cell containing a Science Bowl Question."""

    _q_type_possible = _compile(r"\s*(Short Answer|SA|Multiple Choice|MC)\s*")
    _choices_re = _compile(r"\s*([W|X|Y|Z]\))\s*")
    _answer_re = _compile(r"\s*(ANSWER:)\s*")

    _choices = ("W)", "X)", "Y)", "Z)")

    def __init__(self, cell: _Cell, error_logger: Optional[ErrorLogger] = None):
        self.error_logger = error_logger
        self.cell = cell
        self.state = QuestionFormatterState.Q_START
        self.q_type = None
        self.q_type_run = None
        self.current_choice = 0
        self.choices_para = {}
        self.found_all = False
        self.errors = []

    def _start_handler(self, para: Paragraph):
        q_type_match = self._q_type_possible.match(para.text)
        if q_type_match:
            # the first run of the paragraph should contain the question start
            q_type_run = para.runs[0]
            run_match = self._q_type_possible.match(q_type_run.text)

            if run_match:
                self.q_type = QuestionType.from_string(run_match.group(1))
                run_length = len(q_type_run.text)
                # if the run contains more than the question type, split
                # the run into two
                if run_match.span()[1] < run_length:
                    q_type_run, _ = split_run_at(para, q_type_run, run_match.span()[1])

                q_type_run.text = self.q_type.value
                q_type_run.italic = True
                self.q_type_run = q_type_run

            else:
                # unfortunately, if someone has italicized a single
                # letter in the question start or something, we
                # will have a problem. in this case, highlight
                # the question.
                highlight_cell_text(self.cell, WD_COLOR_INDEX.RED)
                if self.error_logger is not None:
                    self.error_logger.log_error("Couldn't parse question.")
                return None

            # now, the first part of the stem is the second run
            # if it's not left-padded with 4 spaces, make sure it is
            # doing it this way ensures that any other formatting in
            # the stem is preserved (superscripts, subscripts, etc.)
            stem_run = para.runs[1]
            stem_run.text = "".join(["    ", stem_run.text.lstrip()])

            self.state = QuestionFormatterState.STEM_END
        else:
            pass

    def _choice_handler(self, para: Paragraph):

        choice_match = self._choices_re.match(para.text)

        if choice_match:
            # insert a blank paragraph before this one. only
            # do this if we are looking for W), the first choice
            if self.current_choice == 0:
                para.insert_paragraph_before("")

            # the first run should contain the choice
            choice_run = para.runs[0]
            run_match = self._choices_re.match(choice_run.text)
            if run_match:
                # if we matched the wrong choice, replace it with
                # the right choice
                if run_match.group(1) != self._choices[self.current_choice]:
                    choice_run.text = choice_run.text.replace(
                        run_match.group(1), self._choices[self.current_choice], 1
                    )
                # save text and update the choice we're looking for
                self.choices_para[self.current_choice] = para
                self.current_choice += 1
            else:
                # same problem as above, it is possible that someone
                # italicized half of the choice start.
                highlight_cell_text(self.cell, WD_COLOR_INDEX.RED)
                if self.error_logger is not None:
                    self.error_logger.log_error("Couldn't parse question.")
                return None

            # if we just found Z), we're done looking for choices
            if self.current_choice == 4:
                self.current_choice = 0
                self.state = QuestionFormatterState.ANSWER

    def _answer_handler(self, para: Paragraph):
        # if we find a choice line while looking for an answer,
        # the question is probably not Short Answer. highlight
        # the question type to indicate this

        answer_match = self._answer_re.match(para.text)
        if answer_match:
            para.insert_paragraph_before("")

            # set the font to all-caps for every run in the answer line
            for run in para.runs:
                run.text = run.text.upper()

            # for MC questions, additional checks to make sure
            # answer line matches choice
            if self.q_type is QuestionType.MULTIPLE_CHOICE:
                answer_text = para.text.replace("ANSWER: ", "", 1)

                test_choice_re = _compile(r"\s*([W|X|Y|Z]\)?)\s*")
                test_choice_match = test_choice_re.match(answer_text)
                if test_choice_match:
                    # if answer line is a single letter, copy the text of
                    # the choice over to the answer line
                    if test_choice_match.span()[1] == len(answer_text):
                        test_choice = test_choice_match.group(1)
                        if not test_choice.endswith(")"):
                            test_choice += ")"
                        correct_para = self.choices_para[
                            self._choices.index(test_choice)
                        ]
                        para.text = "ANSWER: "
                        for run in correct_para.runs:
                            new_run = para.add_run(run.text.upper())
                            copy_run_formatting(run, new_run)
                        fuse_consecutive_runs(para)
                    # otherwise, check that the answer line text matches the choice.
                    # if it doesn't raise a linting error.
                    else:
                        choice_num = self._choices.index(test_choice_match.group(1))
                        if self.choices_para[choice_num].text.upper() != answer_text:
                            highlight_paragraph_text(para, WD_COLOR_INDEX.YELLOW)
                            if self.error_logger is not None:
                                self.error_logger.log_error(
                                    "Answer line doesn't match choice."
                                )

            # if we made it here, we found everything.
            self.found_all = True

    def parse(self):
        """Parser function. Takes a preprocessed question cell
        and returns a cell containing a properly-formatted
        Science Bowl question.

        Returns
        -------
        _Cell
        """
        for para in self.cell.paragraphs:

            if self.state is QuestionFormatterState.Q_START:
                self._start_handler(para)

            elif self.state is QuestionFormatterState.STEM_END:
                choice_match = self._choices_re.match(para.text)
                answer_match = self._answer_re.match(para.text)

                if choice_match:
                    if self.q_type is QuestionType.SHORT_ANSWER:
                        self.q_type_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        if self.error_logger is not None:
                            self.error_logger.log_error(
                                "Question type is SA, but has choices."
                            )
                        self.q_type = QuestionType.MULTIPLE_CHOICE
                    self.state = QuestionFormatterState.CHOICES
                    self._choice_handler(para)

                elif answer_match:
                    if self.q_type is QuestionType.MULTIPLE_CHOICE:
                        self.q_type_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        if self.error_logger is not None:
                            self.error_logger.log_error(
                                "Question type is MC, but has no choices."
                            )
                        self.q_type = QuestionType.SHORT_ANSWER
                    self.state = QuestionFormatterState.ANSWER
                    self._answer_handler(para)

            elif self.state is QuestionFormatterState.CHOICES:

                self._choice_handler(para)

            elif self.state is QuestionFormatterState.ANSWER:

                self._answer_handler(para)

        if not self.found_all:
            # if we get through the entire cell without finding all parts of
            # of the question, highlight the cell red
            highlight_cell_text(self.cell, WD_COLOR_INDEX.RED)
            if self.error_logger is not None:
                self.error_logger.log_error("Couldn't parse question.")

        return self.cell


def format_tub_cell(cell: _Cell, error_logger: Optional[ErrorLogger] = None) -> _Cell:
    """Formats the TOSS-UP/BONUS cell. Expands shorthand (TU/B/VB) as well.

    Parameters
    ----------
    cell : _Cell

    Returns
    -------
    _Cell
    """
    tub_possible = _compile(r"\s*(TOSS-UP|BONUS|VISUAL BONUS|TU|B|VB)\b")

    tub_match = tub_possible.match(cell.text)

    if tub_match:
        put = TossUpBonus.from_string(tub_match.group(1)).value
        clear_cell(cell)
        tub_run = cell.paragraphs[0].runs[0]
        tub_run.text = put
        tub_run.italic = None
        tub_run.bold = None
        highlight_cell_text(cell, None)

    # if a match can't be found, highlight the cell red
    else:
        highlight_cell_text(cell, WD_COLOR_INDEX.RED)
        if error_logger is not None:
            error_logger.log_error(
                "Question must be a toss-up, bonus, or visual bonus."
            )

    return cell


def format_subject_cell(
    cell: _Cell, error_logger: Optional[ErrorLogger] = None
) -> _Cell:
    """Formats the Subject cell. Expands shorthand as well.

    Parameters
    ----------
    cell : _Cell

    Returns
    -------
    _Cell
    """
    subject_possible = _compile(
        r"\s*(BIOLOGY|B|CHEMISTRY|C|EARTH AND SPACE|ES|ENERGY|EN|MATH|M|PHYSICS|P)\b"
    )
    subject_match = subject_possible.match(cell.text)

    if subject_match:
        put = Subject.from_string(subject_match.group(1)).value
        clear_cell(cell)
        subj_run = cell.paragraphs[0].runs[0]
        subj_run.text = put
        subj_run.italic = None
        subj_run.bold = None
        highlight_cell_text(cell, None)
    # if a match can't be found, highlight the cell red
    else:
        highlight_cell_text(cell, WD_COLOR_INDEX.RED)
        if error_logger is not None:
            error_logger.log_error("Invalid subject.")

    return cell


def format_row(nsb_table_row: _Row, error_logger: Optional[ErrorLogger] = None) -> _Row:
    """Formats the first three cells (TOSS-UP/BONUS,
    SUBJECT, and QUESTION) of a row from one of Jan's tables.

    Parameters
    ----------
    nsb_table_row : _Row
        A row from one of Jan's question tables.

    Returns
    -------
    _Row
    List
    """
    cells_list = nsb_table_row.cells
    if error_logger is None:
        tub_cell = preprocess_cell(cells_list[0])
        subject_cell = preprocess_cell(cells_list[1])
        ques_cell = preprocess_cell(cells_list[2])
    else:
        tub_cell = preprocess_cell(cells_list[0])
        subject_cell = preprocess_cell(cells_list[1])
        ques_cell = preprocess_cell(cells_list[2])
    # make sure first cell says TOSS-UP, BONUS, or VISUAL BONUS and nothing else
    if tub_cell.text.strip() != "":
        format_tub_cell(tub_cell, error_logger)

    # make sure the second cell says one of our subjects and nothing else
    if subject_cell.text.strip() != "":
        format_subject_cell(subject_cell, error_logger)

    # make sure the third cell has a well-formed question
    if ques_cell.text.strip() != "":
        question_formatter = QuestionCellFormatter(ques_cell, error_logger)
        question_formatter.parse()

    return nsb_table_row


def format_table(table_doc: Document, verbosity=True):
    """Formats a Word document containing a Science Bowl question table.

    Parameters
    ----------
    table_doc : Document
    verbosity : bool, default True
    """
    error_logger = ErrorLogger(verbosity)
    # first row is the header row, so we skip it
    for idx, row in enumerate(table_doc.tables[0].rows[1:]):
        error_logger.set_row(idx)
        format_row(row, error_logger)

    if len(error_logger.errors) > 0:
        print(error_logger)
    else:
        print("Found no errors 😄")