import unittest
from pathlib import Path

from nsbtoolbox import tablewriter

from docx.shared import Pt
from docx import Document

data_dir = Path(__file__).parent / "test_data"


class TestInitializeTable(unittest.TestCase):
    def _check_table_characteristics(self, document, nrows):

        font = document.styles["Normal"].font
        self.assertEqual(font.name, "Times New Roman")
        self.assertEqual(font.size, Pt(11))

        self.assertEqual(len(document.tables), 1)
        table = document.tables[0]
        self.assertEqual(len(table.rows), nrows + 1)

        for idx, cell in enumerate(table.rows[0].cells):
            self.assertAlmostEqual(
                cell.width.inches, tablewriter.COL_WIDTHS[idx], places=2
            )

    def test_intialize_table(self):
        _nrows = (0, 90, 150, 180)
        for nrow in _nrows:
            document = tablewriter.initialize_table(nrow)
            self._check_table_characteristics(document, nrow)


class TestPreprocessCell(unittest.TestCase):

    test_data = Document(data_dir / "test_preprocess.docx")
    temp_data = data_dir / "temp" / "temp_preprocess.docx"

    def test_cell_1(self):
        """This cell contains only an uninterrupted run. It shouldn't be changed."""
        cell = self.test_data.tables[0].rows[0].cells[0]

        expected = [["This is a single run of text that is uninterrupted."]]
        test = []
        for para in tablewriter.preprocess_cell(cell).paragraphs:
            test.append([run.text for run in para.runs])

        self.assertEqual(expected, test)

    def test_cell_2(self):
        """This cell contains only an interrupted run. It should be concatenated to
        a single run."""
        cell = self.test_data.tables[0].rows[1].cells[0]

        expected = [["This is a single run of text that has been interrupted."]]
        test = []
        for para in tablewriter.preprocess_cell(cell).paragraphs:
            test.append([run.text for run in para.runs])

        self.assertEqual(expected, test)

    def test_cell_3(self):
        """This cell contains a run that is italicized. It should remain
        separate from the other run."""
        cell = self.test_data.tables[0].rows[2].cells[0]

        expected = [
            [
                "This",
                " is a single run of text that should not be entirely concatenated.",
            ]
        ]
        test = []
        for para in tablewriter.preprocess_cell(cell).paragraphs:
            test.append([run.text for run in para.runs])

        self.assertEqual(expected, test)

    def test_cell_4(self):
        """This cell contains a space that is italicized. Its formatting
        should be stripped."""
        cell = self.test_data.tables[0].rows[3].cells[0]

        expected = [["This is a single run of test that has an italicized space."]]
        test = []
        for para in tablewriter.preprocess_cell(cell).paragraphs:
            test.append([run.text for run in para.runs])

        self.assertEqual(expected, test)

    def test_cell_5(self):
        """This cell contains a broken run with special formatting. It
        should still be fixed."""
        cell = self.test_data.tables[0].rows[4].cells[0]

        expected = [["C", "6", "H", "15", "O", "6"]]
        test = []
        for para in tablewriter.preprocess_cell(cell).paragraphs:
            test.append([run.text for run in para.runs])

        self.assertEqual(expected, test)

    def test_cell_6(self):
        """This cell contains whitespace that should be removed."""
        cell = self.test_data.tables[0].rows[5].cells[0]

        expected = [["This paragraph contains whitespace."]]
        test = []
        for para in tablewriter.preprocess_cell(cell).paragraphs:
            test.append([run.text for run in para.runs])

        self.assertEqual(expected, test)

    def test_save(self):
        self.test_data.save(self.temp_data)
