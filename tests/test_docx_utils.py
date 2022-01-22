import unittest
from pathlib import Path

from docx import Document
from nsb_toolbox import docx_utils

data_dir = Path(__file__).parent / "test_data"


class TestPreprocessCell(unittest.TestCase):

    test_data = Document(data_dir / "test_preprocess.docx")
    temp_data = data_dir / "temp" / "temp_preprocess.docx"

    def test_cell_1(self):
        """This cell contains only an uninterrupted run. It shouldn't be changed."""
        cell = self.test_data.tables[0].rows[0].cells[0]

        expected = [["This is a single run of text that is uninterrupted."]]
        test = []
        for para in docx_utils.preprocess_cell(cell).paragraphs:
            test.append([run.text for run in para.runs])

        self.assertEqual(expected, test)

    def test_cell_2(self):
        """This cell contains only an interrupted run. It should be concatenated to
        a single run."""
        cell = self.test_data.tables[0].rows[1].cells[0]

        expected = [["This is a single run of text that has been interrupted."]]
        test = []
        for para in docx_utils.preprocess_cell(cell).paragraphs:
            test.append([run.text for run in para.runs])

        self.assertEqual(expected, test)

    def test_cell_3(self):
        """This cell contains a run that is italicized. It should remain
        separate from the other run."""
        cell = self.test_data.tables[0].rows[2].cells[0]

        expected = [
            [
                "This",
                " is a single run of text that should not be entirely concatenated.",
            ]
        ]
        test = []
        for para in docx_utils.preprocess_cell(cell).paragraphs:
            test.append([run.text for run in para.runs])

        self.assertEqual(expected, test)

    def test_cell_4(self):
        """This cell contains a space that is italicized. Its formatting
        should be stripped."""
        cell = self.test_data.tables[0].rows[3].cells[0]

        expected = [["This is a single run of test that has an italicized space."]]
        test = []
        for para in docx_utils.preprocess_cell(cell).paragraphs:
            test.append([run.text for run in para.runs])

        self.assertEqual(expected, test)

    def test_cell_5(self):
        """This cell contains a broken run with special formatting. It
        should still be fixed."""
        cell = self.test_data.tables[0].rows[4].cells[0]

        expected = [["C", "6", "H", "15", "O", "6"]]
        test = []
        for para in docx_utils.preprocess_cell(cell).paragraphs:
            test.append([run.text for run in para.runs])

        self.assertEqual(expected, test)

    def test_cell_6(self):
        """This cell contains whitespace that should be removed."""
        cell = self.test_data.tables[0].rows[5].cells[0]

        expected = [["This paragraph contains whitespace."]]
        test = []
        for para in docx_utils.preprocess_cell(cell).paragraphs:
            test.append([run.text for run in para.runs])

        self.assertEqual(expected, test)

    def test_cell_7(self):
        """This cell contains only whitespace."""
        cell = self.test_data.tables[0].rows[6].cells[0]

        expected = [[""]]
        test = []
        for para in docx_utils.preprocess_cell(cell).paragraphs:
            test.append([run.text for run in para.runs])

        self.assertEqual(expected, test)

    def test_cell_8(self):
        cell = self.test_data.tables[0].rows[7].cells[0]

        expected = [
            "This cell contains two soft returns.",
            "It should be split into three paragraphs.",
        ]

        test = [para.text for para in docx_utils.preprocess_cell(cell).paragraphs]
        self.assertEqual(expected, test)

    def test_save(self):
        self.test_data.save(self.temp_data)
