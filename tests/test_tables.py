from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from nsb_toolbox import tables

data_dir = Path(__file__).parent / "test_data"


init_table_sizes = (0, 30, 60)
init_table_subj = (None, "Chemistry", "Earth and Space")
init_table_set = (None, "HSR", "MSN")
init_table_author = (None, "Rishi", "Andrew")


@pytest.fixture(
    scope="module",
    params=zip(init_table_sizes, init_table_subj, init_table_set, init_table_author),
    ids=["Empty", "30 Rows", "60 Rows"],
)
def initialize_params(request):
    return request.param


@pytest.fixture(scope="module")
def initialized_doc_with_table(initialize_params):
    nrows, subj, set_, name = initialize_params
    raw = tables.RawQuestions.make(nrows, subj=subj, set=set_, name=name)
    return raw.document


@pytest.fixture(scope="module")
def table_cells(initialized_doc_with_table):
    return initialized_doc_with_table.tables[0]._cells


class TestInitialize:
    def test_number_of_tables_in_doc(self, initialized_doc_with_table):
        assert len(initialized_doc_with_table.tables) == 1

    def test_numer_of_rows_in_table(
        self, initialized_doc_with_table, initialize_params
    ):
        generated_rows = len(initialized_doc_with_table.tables[0].rows)
        expected_rows = initialize_params[0]
        assert generated_rows == expected_rows + 1

    def test_cell_widths(self, table_cells):

        for idx, cell in enumerate(table_cells):
            row_idx, col_idx = divmod(idx, 13)

            assert pytest.approx(cell.width.inches, 0.001) == tables.COL_WIDTHS[col_idx]

    def test_cell_optional_parameters(self, table_cells, initialize_params):

        _, subj, set_, author = initialize_params

        for idx, cell in enumerate(table_cells):
            row_idx, col_idx = divmod(idx, 13)

            if col_idx == 1 and row_idx > 0 and subj:
                assert cell.text == subj

            if col_idx == 5 and row_idx > 0 and set_:
                assert cell.text == set_

            if col_idx == 8 and row_idx > 0 and author:
                assert cell.text == author


@pytest.fixture
def tub_test_doc():
    return Document(data_dir / "test_TUB.docx")


@pytest.mark.parametrize("row_idx", [0, 1, 2], ids=["TOSS-UP", "BONUS", "VISUAL BONUS"])
class TestTUBCellFormatter:
    def test_formatted_text(self, tub_test_doc, row_idx):
        expected_text = ["TOSS-UP", "BONUS", "VISUAL BONUS"]
        row = tub_test_doc.tables[0].rows[row_idx]
        for cell in row.cells:
            formatted_cell = tables.TuBCellFormatter().format(cell)
            assert formatted_cell.text == expected_text[row_idx]

    def test_cell_has_single_paragraph(
        self,
        tub_test_doc,
        row_idx,
    ):
        row = tub_test_doc.tables[0].rows[row_idx]
        for cell in row.cells:
            formatted_cell = tables.TuBCellFormatter().format(cell)
            assert len(formatted_cell.paragraphs) == 1

    def test_paragraph_contains_one_run_with_normal_text(self, tub_test_doc, row_idx):
        row = tub_test_doc.tables[0].rows[row_idx]
        for cell in row.cells:
            formatted_cell = tables.TuBCellFormatter().format(cell)
            cell_runs = formatted_cell.paragraphs[0].runs
            assert len(cell_runs) == 1
            assert cell_runs[0].font.italic is None
            assert cell_runs[0].font.bold is None


class TestTUBCellFormatterErrors:
    def test_unrecognizable_cell_text_is_unchanged(self, tub_test_doc):
        row = tub_test_doc.tables[0].rows[3]
        for cell in row.cells:
            prior_text = cell.text
            formatted_cell = tables.TuBCellFormatter().format(cell)
            after_text = formatted_cell.text
            assert prior_text == after_text

    def test_unrecognizable_cell_is_highlighted(self, tub_test_doc):
        row = tub_test_doc.tables[0].rows[3]
        for cell in row.cells:
            formatted_cell = tables.TuBCellFormatter().format(cell)
            run = formatted_cell.paragraphs[0].runs[0]
            assert run.font.highlight_color == WD_COLOR_INDEX.RED


@pytest.fixture
def format_difficulty_rows():
    test_data = Document(data_dir / "test_LOD.docx")
    return test_data.tables[0].rows


@pytest.mark.parametrize("cell_idx", [0, 1, 2, 3])
class TestDifficultyFormatter:
    def test_expected_test(self, format_difficulty_rows, cell_idx):

        EXPECTED_TEXT = (
            "1",
            "2",
            "3",
            "",
        )
        cell = format_difficulty_rows[0].cells[cell_idx]
        formatter = tables.DifficultyFormatter()
        assert formatter.format(cell).text == EXPECTED_TEXT[cell_idx]

    def test_difficulty_errors(self, format_difficulty_rows, cell_idx):

        cell = format_difficulty_rows[1].cells[cell_idx]
        formatter = tables.DifficultyFormatter()
        test_cell = formatter.format(cell)
        assert (
            test_cell.paragraphs[0].runs[0].font.highlight_color == WD_COLOR_INDEX.RED
        )


@pytest.fixture
def format_subject_rows():
    test_data = Document(data_dir / "test_subject.docx")
    return test_data.tables[0].rows


@pytest.mark.parametrize(
    "row_idx",
    [0, 1, 2, 3, 4, 5],
    ids=[
        "Biology",
        "Chemistry",
        "Physics",
        "Earth and Space",
        "Math",
        "Energy",
    ],
)
@pytest.mark.parametrize(
    "cell_idx",
    [0, 1, 2, 3],
    ids=[
        "Full Correct",
        "Abbreviated",
        "Abbreviated Wrong Format",
        "Full Wrong Format",
    ],
)
class TestFormatSubject:
    def test_expected_text(self, format_subject_rows, row_idx, cell_idx):
        EXPECTED_TEXT = (
            "Biology",
            "Chemistry",
            "Physics",
            "Earth and Space",
            "Math",
            "Energy",
        )
        cell = format_subject_rows[row_idx].cells[cell_idx]
        formatter = tables.SubjectCellFormatter()
        assert formatter.format(cell).text == EXPECTED_TEXT[row_idx]

    def test_cell_formatting(self, format_subject_rows, row_idx, cell_idx):
        cell = format_subject_rows[row_idx].cells[cell_idx]
        formatter = tables.SubjectCellFormatter()

        formatted_cell = formatter.format(cell)
        assert len(formatted_cell.paragraphs) == 1
        assert len(formatted_cell.paragraphs[0].runs) == 1
        assert formatted_cell.paragraphs[0].runs[0].font.italic is None
        assert formatted_cell.paragraphs[0].runs[0].font.bold is None


@pytest.mark.parametrize(
    "cell_idx",
    [0, 1, 2, 3],
    ids=[
        "Biologychemistry",
        "Multiple Choice",
        "_Life Science_",
        "_Physical Science_",
    ],
)
class TestFormatSubjectErrors:
    def test_errors(self, format_subject_rows, cell_idx):
        cell = format_subject_rows[6].cells[cell_idx]

        prior_text = cell.text
        sformatter = tables.SubjectCellFormatter()
        test_run = sformatter.format(cell).paragraphs[0].runs[0]
        after_text = cell.text
        assert prior_text == after_text
        assert test_run.font.highlight_color == WD_COLOR_INDEX.RED


@pytest.fixture
def format_question_rows():
    test_data = Document(data_dir / "test_question_parser.docx")
    return test_data.tables[0].rows


@pytest.mark.parametrize(
    "cell_idx",
    [0, 1, 2, 3, 4, 5],
    ids=[
        "Properly formatted",
        "Extra whitespace before stem",
        "Extra whitespace paragraphs",
        "Abbreviated",
        "Linebreak",
        "Linebreak with extra whitespace",
    ],
)
class TestQuestionFormat:
    def _extract_cell_text(self, cell):
        ret = []
        for para in cell.paragraphs:
            if para.runs == []:
                ret.append([""])
            else:
                ret.append([run.text for run in para.runs])

        return ret

    @pytest.mark.parametrize(
        "force_capitalize", [True, False], ids=["+capitalize", "-capitalize"]
    )
    def test_short_answer(self, format_question_rows, cell_idx, force_capitalize):
        """This makes sure that recognizable Short Answer questions
        are properly formatted."""
        if force_capitalize:
            expected = [
                ["Short Answer", "    This is a well-formatted question."],
                [""],
                ["ANSWER: IT SHOULD BE UNCHANGED"],
            ]
        else:
            expected = [
                ["Short Answer", "    This is a well-formatted question."],
                [""],
                ["ANSWER: it should be unchanged"],
            ]

        cell = format_question_rows[0].cells[cell_idx]

        q_parser = tables.QuestionCellFormatter(force_capitalize=force_capitalize)
        test_text = self._extract_cell_text(q_parser.preprocess_format(cell))
        assert test_text == expected
        assert cell.paragraphs[-1].runs[0].font.highlight_color is None

    @pytest.mark.parametrize(
        "force_capitalize", [True, False], ids=["+capitalize", "-capitalize"]
    )
    def test_multiple_choice_answer_given(
        self, format_question_rows, cell_idx, force_capitalize
    ):
        """This makes sure that recognizable Short Answer questions
        are properly formatted."""
        if force_capitalize:
            expected = [
                ["Multiple Choice", "    This is a well-formatted question."],
                ["W) This is the W) choice"],
                ["X) This is the X) choice"],
                ["Y) This is the Y) choice"],
                ["Z) This is the Z) choice"],
                [""],
                ["ANSWER: W) THIS IS THE W) CHOICE"],
            ]
        else:
            expected = [
                ["Multiple Choice", "    This is a well-formatted question."],
                ["W) This is the W) choice"],
                ["X) This is the X) choice"],
                ["Y) This is the Y) choice"],
                ["Z) This is the Z) choice"],
                [""],
                ["ANSWER: W) this is the w) choice"],
            ]

        cell = format_question_rows[1].cells[cell_idx]

        q_parser = tables.QuestionCellFormatter(force_capitalize=force_capitalize)
        test_text = self._extract_cell_text(q_parser.preprocess_format(cell))
        assert test_text == expected
        assert cell.paragraphs[-1].runs[0].font.highlight_color is None

    def test_multiple_choice_fill_answer(self, format_question_rows, cell_idx):
        """This makes sure that recognizable Short Answer questions
        are properly formatted."""
        expected = [
            ["Multiple Choice", "    This is a well-formatted question."],
            ["W) This is the W) choice"],
            ["X) This is the X) choice"],
            ["Y) This is the Y) choice"],
            ["Z) This is the Z) choice"],
            [""],
            ["ANSWER: W) THIS IS THE W) CHOICE"],
        ]

        cell = format_question_rows[2].cells[cell_idx]

        q_parser = tables.QuestionCellFormatter()
        test_text = self._extract_cell_text(q_parser.preprocess_format(cell))
        assert test_text == expected
        assert cell.paragraphs[-1].runs[0].font.highlight_color is None

    @pytest.mark.parametrize(
        "force_capitalize", [True, False], ids=["+capitalize", "-capitalize"]
    )
    def test_SA_intentional_line_break_handling(
        self, format_question_rows, cell_idx, force_capitalize
    ):
        """This makes sure that recognizable Short Answer questions
        are properly formatted."""
        if force_capitalize:
            expected = [
                ["Short Answer", "    This is a well-formatted question that is"],
                ["split across multiple lines on purpose"],
                [""],
                ["ANSWER: IT SHOULD BE UNCHANGED"],
            ]
        else:
            expected = [
                ["Short Answer", "    This is a well-formatted question that is"],
                ["split across multiple lines on purpose"],
                [""],
                ["ANSWER: it should be unchanged"],
            ]

        cell = format_question_rows[3].cells[cell_idx]

        q_parser = tables.QuestionCellFormatter(force_capitalize=force_capitalize)
        test_text = self._extract_cell_text(q_parser.preprocess_format(cell))
        assert test_text == expected
        assert cell.paragraphs[-1].runs[0].font.highlight_color is None

    @pytest.mark.parametrize(
        "force_capitalize", [True, False], ids=["+capitalize", "-capitalize"]
    )
    def test_MC_intentional_line_break_handling(
        self, format_question_rows, cell_idx, force_capitalize
    ):
        """This makes sure that recognizable Short Answer questions
        are properly formatted."""
        if force_capitalize:
            expected = [
                ["Multiple Choice", "    This is a well-formatted question."],
                ["W) This is the W) choice"],
                ["X) This is the X) choice"],
                ["Y) This is the Y) choice"],
                ["Z) This is the Z) choice and it is"],
                ["split across multiple lines on purpose"],
                [""],
                ["ANSWER: W) THIS IS THE W) CHOICE"],
            ]
        else:
            expected = [
                ["Multiple Choice", "    This is a well-formatted question."],
                ["W) This is the W) choice"],
                ["X) This is the X) choice"],
                ["Y) This is the Y) choice"],
                ["Z) This is the Z) choice and it is"],
                ["split across multiple lines on purpose"],
                [""],
                ["ANSWER: W) this is the w) choice"],
            ]

        cell = format_question_rows[4].cells[cell_idx]

        q_parser = tables.QuestionCellFormatter(force_capitalize=force_capitalize)
        test_text = self._extract_cell_text(q_parser.preprocess_format(cell))
        assert test_text == expected
        assert cell.paragraphs[-1].runs[0].font.highlight_color is None


class TestQuestionFormatterErrors:
    @pytest.mark.parametrize(
        "cell_idx",
        [0, 1],
        ids=["Incorrectly labeled MC", "Incorrectly labeled SA"],
    )
    def test_question_type_warning(self, format_question_rows, cell_idx):
        """Tests that mislabeled question types get warnings."""
        cell = format_question_rows[5].cells[cell_idx]

        q_parser = tables.QuestionCellFormatter()
        test_cell = q_parser.preprocess_format(cell)
        # check that qtype run is highlighted
        assert (
            test_cell.paragraphs[0].runs[0].font.highlight_color
            == WD_COLOR_INDEX.YELLOW
        )
        # check that other runs aren't highlighted
        assert cell.paragraphs[-1].runs[0].font.highlight_color is None

    @pytest.mark.parametrize(
        "cell_idx",
        [0, 1, 2, 3, 4, 5],
        ids=[
            "Properly formatted",
            "Extra whitespace before stem",
            "Extra whitespace paragraphs",
            "Abbreviated",
            "Linebreak",
            "Linebreak with extra whitespace",
        ],
    )
    def test_answer_line_warning(self, format_question_rows, cell_idx):
        """Tests that answer lines that don't match choices in
        MC questions get warnings."""
        cell = format_question_rows[6].cells[cell_idx]

        q_parser = tables.QuestionCellFormatter()
        test_cell = q_parser.preprocess_format(cell)
        assert (
            test_cell.paragraphs[-1].runs[0].font.highlight_color
            == WD_COLOR_INDEX.YELLOW
        )

    @pytest.mark.parametrize(
        "cell_idx",
        [0, 1, 2, 3, 4, 5],
        ids=[
            "MC without 4 choices",
            "Question type split across multiple runs",
            "Missing stem",
            "Choice split across multiple runs",
            "Answer choice doesn't match available choices",
            "Answer choice is two letters",
        ],
    )
    def test_malformed_question_error(self, format_question_rows, cell_idx):
        cell = format_question_rows[7].cells[cell_idx]

        q_parser = tables.QuestionCellFormatter()
        test_cell = q_parser.preprocess_format(cell)
        assert (
            test_cell.paragraphs[-1].runs[0].font.highlight_color == WD_COLOR_INDEX.RED
        )


@pytest.fixture(scope="module")
def format_test_table():
    doc = Document(data_dir / "test_format.docx")
    tables.RawQuestions(doc).format(
        force_capitalize=False,
    )
    return doc.tables[0]


@pytest.mark.parametrize(
    "cell_idx",
    [0, 1, 2, 3, 4, 5],
    ids={
        "Header",
        "Question 1",
        "Question 2",
        "Question 3",
        "Question 4",
        "Question 5",
    },
)
class TestFormat:
    def test_font_name(self, format_test_table, cell_idx):
        assert format_test_table.part.styles["Normal"].font.name == "Times New Roman"

    def test_font_size(self, format_test_table, cell_idx):
        assert format_test_table.part.styles["Normal"].font.size == Pt(12)

    def test_tub_col(self, format_test_table, cell_idx):
        expected = ["TUB", "TOSS-UP", "BONUS", "TOSS-UP", "BONUS", "TOSS-UP"]
        expected_text = expected[cell_idx]

        assert format_test_table.columns[0].cells[cell_idx].text == expected_text

    def test_subj_col(self, format_test_table, cell_idx):
        expected = [
            "Subj",
            "Chemistry",
            "Biology",
            "Earth and Space",
            "Energy",
            "Physics",
        ]
        expected_text = expected[cell_idx]

        assert format_test_table.columns[1].cells[cell_idx].text == expected_text

    def test_ques_col(self, format_test_table, cell_idx):
        expected = [
            "Ques",
            "Short Answer    Question\n\nANSWER: answer",
            "Short Answer    Question\n\nANSWER: ANSWER",
            "Multiple Choice    Question\nW) w\nX) x\nY) y\nZ) z\n\nANSWER: Z) z",
            "Multiple Choice    Question\nW) w\nX) x\nY) y\nZ) z\n\nANSWER: Z) Z",
            "Multiple Choice    Question\nW) w\nX) x\nY) y\nZ) z\n\nANSWER: Z) Z",
        ]

        expected_text = expected[cell_idx]

        assert format_test_table.columns[2].cells[cell_idx].text == expected_text

    def test_lod_col(self, format_test_table, cell_idx):
        expected = [
            "LOD",
            "1",
            "2",
            "3",
            "4",
            "",
        ]
        expected_text = expected[cell_idx]

        assert format_test_table.columns[3].cells[cell_idx].text == expected_text
