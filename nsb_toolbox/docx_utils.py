from typing import Generator
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_COLOR_INDEX

from copy import deepcopy


def preprocess_cell(cell: _Cell) -> _Cell:
    """Multipass cleaning function for table cells.

    Parameters
    ----------
    cell : _Cell

    Returns
    -------
    _Cell
    """
    highlight_cell_text(cell, None)

    if cell.text.strip() == "":
        clear_cell(cell)

    else:
        # this pass replaces all soft returns with hard returns
        for para in cell.paragraphs:
            split_soft_returns(para)

        for para in cell.paragraphs:

            # this pass coerces the font of any whitespace-only runs to
            # the document style
            for run in para.runs:
                # replace non-breaking spaces with regular spaces
                run.text = run.text.replace("\xa0", " ")
                # if there are empty runs, delete them
                if run.text == "":
                    delete_run(run)
                # if there is a weirdly formatted run that is only whitespace,
                # strip their formatting
                elif run.text.strip() == "":
                    run.font.italic = (
                        run.font.bold
                    ) = (
                        run.font.all_caps
                    ) = (
                        run.font.subscript
                    ) = run.font.superscript = run.font.underline = None

            # this pass combines runs that have the same font properties
            # editing an XML file splits a run, so this is necessary
            fuse_consecutive_runs(para)

            # finally, delete any left padding or right padding for cells
            # containing text delete paragraphs that are empty
            if para.text.strip() == "":
                delete_paragraph(para)
            else:
                strip_padding_whitespace(para, leading=True)
                strip_padding_whitespace(para, leading=False)

    return cell


def strip_padding_whitespace(para: Paragraph, leading=True):
    """Removes any empty runs at the beginning or end of a paragraph
    and makes sure the first (or last) run with text has no left
    (or right) padding.

    Parameters
    ----------
    para : Paragraph
    leading : bool, optional
        If true, strips left padding, otherwise strips right padding,
        by default True
    """
    if leading:
        target_idx = 0

    else:
        target_idx = -1

    target_run = para.runs[target_idx]

    while not target_run.text.strip():
        delete_run(target_run)
        target_run = para.runs[target_idx]

    if leading:
        target_run.text = target_run.text.lstrip()
    else:
        target_run.text = target_run.text.rstrip()


def split_soft_returns(para: Paragraph) -> Paragraph:
    """Find and replace function for turning soft returns into hard returns.

    In OXML, soft returns are the same paragraph, while hard returns start
    a new paragraph. Don't want to deal with two different kinds of line
    breaks, so force them all the be the same kind.

    Parameters
    ----------
    para : Paragraph

    Returns
    -------
    Paragraph
    """
    idx = 0
    for run in para.runs:
        while (newline_loc := run.text.find("\n")) != -1:
            _, run_2 = split_run_at(run, newline_loc)
            run_2.text = run_2.text[1:]  # remove the newline char

            # move all runs prior to this one to a new paragraph above
            new_para = para.insert_paragraph_before(" ")
            new_para._p[:] = para._p[: idx + 1]

            idx = 0
        idx += 1
    return para


def fuse_consecutive_runs(para: Paragraph) -> Paragraph:
    """Compares every run in a paragraph with the next. If they
    have the same formatting, they are combined into a single
    run.

    Parameters
    ----------
    para : Paragraph

    Returns
    -------
    Paragraph
    """
    for run_1, run_2 in zip(para.runs[:-1], para.runs[1:]):
        if compare_run_styles(run_1, run_2):
            run_2.text = run_1.text + run_2.text
            delete_run(run_1)
    return para


def compare_run_styles(run_1: Run, run_2: Run) -> bool:
    """Nonexhaustively compares two runs to check if they have
    the same font. Science Bowl only uses italic, bold, all caps,
    superscript, subscript, and underline, so only those are
    compared.

    Parameters
    ----------
    run_1, run_2 : Run

    Returns
    -------
    bool
    """
    font_1 = run_1.font
    font_2 = run_2.font

    return (
        (font_1.italic == font_2.italic)
        and (font_1.bold == font_2.bold)
        and (font_1.all_caps == font_2.all_caps)
        and (font_1.superscript == font_2.superscript)
        and (font_1.subscript == font_2.subscript)
        and (font_1.underline == font_1.underline)
    )


def split_run_at(run: Run, split_at: int):
    """Splits a run at a specified index.

    Parameters
    ----------
    run : Run
    split_at : int
        Index of split location in the run.

    Returns
    -------
    list of runs
    """
    txt = run.text
    split_at %= len(txt)
    left, right = [txt[:split_at], txt[split_at:]]

    run_2 = deepcopy(run._r)

    run.text, run_2.text = left, right

    # move second run to be after first run
    run._r.addnext(run_2)

    return [run, run_2]


def move_runs_to_end_of_para(para_from: Paragraph, para_to: Paragraph) -> None:
    """Moves every run from para_from to the end of para_to.

    Parameters
    ----------
    para_from : Paragraph
    para_to : Paragraph
    """
    para_to.extend(para_from)
    para_from.getparent().remove(para_from)


def shade_cell(cell, shade: str) -> None:
    """Shades a list of cells in-place with a hex color value.

    Parameters
    ----------
    cells : Iterable
        Cells to shade
    shade : str
        Hexadecimal color value
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcVAlign = OxmlElement("w:shd")
    tcVAlign.set(qn("w:fill"), shade)
    tcPr.append(tcVAlign)


def delete_paragraph(paragraph: Paragraph) -> None:
    """Deletes a paragraph."""
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def delete_run(run: Run) -> None:
    """Deletes a run."""
    p = run._r.getparent()
    p.remove(run._r)


def clear_cell(cell: _Cell) -> None:
    """Deletes every paragraph in a cell except for the first, and makes the first
    paragraph contain only an empty run of text. Removes shading (and all other)
    element attributes, as well.

    Parameters
    ----------
    cell : _Cell
    """
    cell._tc.clear()
    cell.add_paragraph("").add_run("")


def highlight_cell_text(cell: _Cell, color: WD_COLOR_INDEX) -> None:
    """Highlights all the text in a cell a given color. Used for
    providing linter warnings.

    Parameters
    ----------
    cell : _Cell
    color : WD_COLOR_INDEX
    """
    for paragraph in cell.paragraphs:
        highlight_paragraph_text(paragraph, color)


def highlight_paragraph_text(para: Paragraph, color: WD_COLOR_INDEX) -> None:
    """Highlights every run in a paragraph a given color.

    Parameters
    ----------
    para : Paragraph
    color : WD_COLOR_INDEX
    """
    for run in para.runs:
        run.font.highlight_color = color


def capitalize_paragraph(para: Paragraph) -> None:
    """Capitalizes all text in a paragraph.

    Parameters
    ----------
    para : Paragraph
    """
    for run in para.runs:
        run.text = run.text.upper()


def column_indexer(
    col_num: int, total_cells: int, col_count: int, skip_header: bool = True
) -> Generator[int, None, None]:
    """Convenience function to build iterators over columns in a table.

    Parameters
    ----------
    col_num : int
        Column number to iterate over, indexing starts at 0.
    total_cells : int
        Total number of cells in the table.
    col_count : int
        Number of columns in the table.
    skip_header : bool, optional
        If true, the iterator will skip the first instance, by default True

    Yields
    ------
    Generator[int, None, None]
        range generator that yields cell indexes for the column of interest.
    """
    if skip_header:
        return range(col_num + col_count, total_cells, col_count)
    else:
        return range(col_num, total_cells, col_count)
