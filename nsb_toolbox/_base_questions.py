from functools import cached_property
from pathlib import Path
from typing import Dict, List, Union

import docx.document
import docx.table
import numpy as np
from typing_extensions import Self

from nsb_toolbox.importers import load_doc

from .classes import QuestionType, TossUpBonus
from .docx_utils import column_indexer


class BaseScienceBowlQuestions:
    """Represents an edited set of Science Bowl questions from a single subject
    with level of difficulty and subcategory assigned by the SME.

    Parameters
    ----------
    document : docx.document.Document
        Word document of Science Bowl questions wrapped by this class.

    Raises
    ------
    ValueError
        Raised if there are formatting issues in the document.

    Attributes
    ----------
    tubs : np.ndarray[str]
    difficulties : np.ndarray[int]
    qtypes : np.ndarray[str]
    subcategories : np.ndarray[str]
    sets : List[docx.table._Cell]
    rounds : List[docx.table._Cell]
    qletters : List[docx.table._Cell]

    Methods
    -------
    assign(question_spec: ParsedQuestionSpec)
        Assigns the questions to a specification via a linear sum assignment.
    """

    def __init__(self, document: docx.document.Document) -> Self:

        self.document = document

        self._cells = self.document.tables[0]._cells
        self._col_count = self.document.tables[0]._column_count

    def save(self, path: Union[Path, str]):
        """Saves the wrapped document to path.

        Parameters
        ----------
        path : Union[Path, str]
        """
        self.document.save(path)

    @classmethod
    def from_docx_path(cls, path: Union[Path, str], dry_run: bool = False) -> Self:
        """Generates a class instance from a path to a docx file.

        Parameters
        ----------
        path : Union[Path, str]

        Returns
        -------
        EditedQuestions
        """
        doc = load_doc(path)
        return cls(doc)

    @cached_property
    def tubs(self) -> np.ndarray:
        return np.array(
            [
                TossUpBonus(self._cells[i].text).value
                for i in column_indexer(0, len(self._cells), self._col_count)
            ],
            dtype="<U20",
        )

    @cached_property
    def difficulties(self) -> np.ndarray:
        return np.array(
            [
                int(self._cells[i].text or -1)
                for i in column_indexer(3, len(self._cells), self._col_count)
            ]
        )

    @cached_property
    def qtypes(self) -> np.ndarray:
        return np.array(
            [
                QuestionType(qtype).value
                if (qtype := self._cells[i].paragraphs[0].runs[0].text)
                else ""
                for i in column_indexer(2, len(self._cells), self._col_count)
            ],
            dtype="<U20",
        )

    @property
    def qtype_stats(self) -> Dict[str, int]:
        return {
            val: count
            for val, count in zip(*np.unique(self.qtypes, return_counts=True))
            if val
        }

    @cached_property
    def subcategories(self) -> np.ndarray:
        return np.array(
            [
                self._cells[i].text
                for i in column_indexer(12, len(self._cells), self._col_count)
            ],
            dtype="<U20",
        )

    @cached_property
    def writers(self) -> np.ndarray:
        return np.array(
            [
                self._cells[i].text
                for i in column_indexer(8, len(self._cells), self._col_count)
            ],
            dtype="<U100",
        )

    @cached_property
    def stats(self) -> Dict[str, int]:
        raw_values = np.array(
            [
                f"{set_.text:<10}{difficulty:^}{tub:>12}"
                for set_, difficulty, tub in zip(
                    self.sets, self.difficulties, self.tubs
                )
            ]
        )
        return {
            val: count for val, count in zip(*np.unique(raw_values, return_counts=True))
        }

    @property
    def sets(self) -> List[docx.table._Cell]:
        return [
            self._cells[i] for i in column_indexer(5, len(self._cells), self._col_count)
        ]

    @property
    def rounds(self) -> List[docx.table._Cell]:
        return [
            self._cells[i] for i in column_indexer(6, len(self._cells), self._col_count)
        ]

    @property
    def qletters(self) -> List[docx.table._Cell]:
        return [
            self._cells[i] for i in column_indexer(7, len(self._cells), self._col_count)
        ]
