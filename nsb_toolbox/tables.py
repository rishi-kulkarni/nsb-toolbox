import logging
import re
from copy import deepcopy
from enum import Enum
from functools import cached_property, partial
from typing import Dict, Iterable, Optional, Union
from typing_extensions import Self

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches, Pt
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
import numpy as np

from ._base_questions import BaseScienceBowlQuestions

from .classes import (
    QuestionType,
    Subject,
    TossUpBonus,
    RowContextFilter,
    FormatErrorsFormatter,
)
from .docx_utils import (
    capitalize_paragraph,
    clear_cell,
    column_indexer,
    fuse_consecutive_runs,
    highlight_cell_text,
    highlight_paragraph_text,
    move_runs_to_end_of_para,
    preprocess_cell,
    shade_cell,
    split_run_at,
)


logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
row_filter = RowContextFilter()
logger.addFilter(row_filter)

ch = logging.StreamHandler()
ch.setLevel(logging.INFO)
msg_formatter = FormatErrorsFormatter()
ch.setFormatter(msg_formatter)

logger.addHandler(ch)

COL_WIDTHS = (
    0.72,
    0.61,
    4.71,
    0.58,
    0.68,
    0.68,
    0.68,
    0.68,
    0.81,
    0.64,
    4.71,
    1.51,
    1.44,
)

COL_MAPPING = {
    "TUB": 0,
    "Subj": 1,
    "Ques": 2,
    "LOD": 3,
    "LOD-A": 4,
    "Set": 5,
    "Rd": 6,
    "Q Letter": 7,
    "Author": 8,
    "ID": 9,
    "Source": 10,
    "Comments": 11,
    "Subcat": 12,
}

COL_COLORS = {
    "TUB": None,
    "Subj": None,
    "Ques": None,
    "LOD": "#FFCC99",
    "LOD-A": None,
    "Set": "#e5dfec",
    "Rd": "#daeef3",
    "Q Letter": "#daeef3",
    "Author": None,
    "ID": None,
    "Source": None,
    "Comments": None,
    "Subcat": None,
}

TUB_RE = re.compile(r"\s*(TOSS-UP|BONUS|VISUAL BONUS|TU|B|VB)\b", re.IGNORECASE)
SUBJECT_RE = re.compile(
    r"\s*(BIOLOGY|B|CHEMISTRY|C|EARTH AND SPACE|ES|ENERGY|EN|MATH|M|PHYSICS|P)\b",
    re.IGNORECASE,
)
Q_TYPE_RE = re.compile(r"\s*(Short Answer|SA|Multiple Choice|MC)\s*", re.IGNORECASE)
CHOICES_RE = re.compile(r"\s*([W|X|Y|Z]\))\s*", re.IGNORECASE)
TEST_CHOICE_RE = re.compile(r"(?:ANSWER:)?\s*([W|X|Y|Z])(?:\)?$|\).+)", re.IGNORECASE)
ANSWER_RE = re.compile(r"\s*(ANSWER:)\s*", re.IGNORECASE)

CHOICES = ("W)", "X)", "Y)", "Z)")


class RawQuestions(BaseScienceBowlQuestions):
    @classmethod
    def make(
        cls,
        nrows: int = 0,
        name: Optional[str] = None,
        subj: Optional[str] = None,
        set: Optional[str] = None,
    ) -> Self:
        """Initializes a class instance file containing a newly-initialized
        raw question table.

        Parameters
        ----------
        nrows : int, optional
            Number of extra rows to append to the table, by default 0

        name : str, optional
            Name of author. If not none, fills the author column of the table.

        subj: str, optional
            Subject. If not none, fills the subject column of the table.

        set: str, optional
            Set. If not none, fills the set column of the table.

        Returns
        -------
        Self
        """
        document = Document()

        table = document.add_table(rows=1 + nrows, cols=13)

        table.style = "Table Grid"
        table.autofit = False
        table.allow_autofit = False

        _cells = table._cells
        _col_count = table._column_count

        col_iter = partial(
            column_indexer,
            total_cells=len(_cells),
            col_count=_col_count,
            skip_header=False,
        )

        for col_name, col_idx in COL_MAPPING.items():
            for cell_idx in col_iter(col_idx):
                cell = _cells[cell_idx]
                cell.width = Inches(COL_WIDTHS[col_idx])

                if cell_idx < _col_count:
                    cell.paragraphs[0].add_run(col_name)

                elif col_name == "Subj" and subj is not None:
                    cell.paragraphs[0].text = subj

                elif col_name == "Set":
                    if set is not None:
                        cell.paragraphs[0].text = set

                elif col_name == "Author" and name is not None:
                    cell.paragraphs[0].text = name

        # ques header is italicized
        ques_run = table.cell(0, COL_MAPPING["Ques"]).paragraphs[0].runs[0]
        ques_run.italic = True

        return cls(document)

    def format(
        self,
        force_capitalize: bool = False,
        verbose: bool = True,
        line_after_stem: bool = False,
    ):
        """Formats a Word document containing a Science Bowl question table.

        Specifically, this function makes sure the columns fit the following
        criteria:

        TUB: Contains only TOSS-UP or BONUS
        Subj: Contains only one of the valid subject areas
        Ques: Contains a properly-formatted Short Answer or Multiple Choice question
        LOD: Contains only an integer or is blank.
        Set: Contains no extra whitespace
        Author: Contains no extra whitespace
        Subcat: Contains no extra whitespace

        Parameters
        ----------
        force_capitalize : bool, default True
            If True, all answer lines will be capitalized
        """

        cols_to_format = (
            "TUB",
            "Subj",
            "Ques",
            "LOD",
            "Set",
            "Rd",
            "Q Letter",
            "Author",
            "Subcat",
        )

        FORMATTERS = {
            "TUB": TuBCellFormatter(),
            "Subj": SubjectCellFormatter(),
            "Ques": QuestionCellFormatter(
                force_capitalize=force_capitalize, line_after_stem=line_after_stem
            ),
            "LOD": DifficultyFormatter(),
            "Set": SetFormatter(),
            "Rd": RdFormatter(),
            "Q Letter": QLetterFormatter(),
        }

        _cells = self.document.tables[0]._cells
        _col_count = self.document.tables[0]._column_count

        font = self.document.styles["Normal"].font
        font.name = "Times New Roman"
        font.size = Pt(12)

        col_iter = partial(
            column_indexer,
            total_cells=len(_cells),
            col_count=_col_count,
            skip_header=True,
        )

        for col_name in cols_to_format:
            formatter = FORMATTERS.get(col_name, CellFormatter())
            formatter.preprocess_format_column(
                _cells[i] for i in col_iter(COL_MAPPING[col_name])
            )
        if verbose:
            if row_filter._num_records == 0:
                print("Found no errors 😄")
            self.print_stats()

        return self

    def print_stats(self):
        print("\nStatistics")
        print(f"{'-':->34}")
        print(f"{'Set':<9}{'LOD':^}{'TUB':>10}{'Have':>11}")
        print(f"{'-':->34}")

        for key in sorted(self.stats.keys()):
            have = self.stats.get(key, 0)
            print(f"{key}{have:>10}")

        print(f"{'-':->34}")
        for key in sorted(self.qtype_stats.keys()):
            have = self.qtype_stats.get(key, 0)
            print(f"{key :<15}{have :>18}")

    @cached_property
    def tubs(self) -> np.ndarray:
        """Overridden because TUB might be malformed."""
        return np.array(
            [
                TossUpBonus(self._cells[i].text).value
                if self._cells[i].text in ("TOSS-UP", "BONUS", "VISUAL BONUS")
                else "ERROR"
                for i in column_indexer(0, len(self._cells), self._col_count)
            ],
            dtype="<U20",
        )

    @cached_property
    def difficulties(self) -> np.ndarray:
        """Overridden because difficulties may be written in LOD-A column."""
        return np.array(
            [
                int(self._cells[lod].text or self._cells[lod_a].text or -1)
                for lod, lod_a in zip(
                    column_indexer(3, len(self._cells), self._col_count),
                    column_indexer(4, len(self._cells), self._col_count),
                )
            ]
        )


class CellFormatter:
    """Base class that ensures formatters are standardized."""

    def format(self, cell: _Cell) -> _Cell:
        """All CellFormatters have a format function."""
        return cell

    def preprocess_format(self, cell: _Cell) -> _Cell:
        """Convenience function to preprocess and format a cell."""
        cell = preprocess_cell(cell)
        if hasattr(self, "color"):
            shade_cell(cell, self.color)
        if cell.text.strip() != "":
            return self.format(cell)

    def preprocess_format_column(self, cells: Iterable[_Cell]) -> _Cell:
        """Convenience function to preprocess and format an entire column."""
        for idx, cell in enumerate(cells):
            row_filter.curr_row = idx + 1
            self.preprocess_format(cell)


class QuestionParserException(Exception):
    pass


class QuestionFormatterState(Enum):
    Q_START = 0
    STEM_END = 1
    CHOICES = 2
    ANSWER = 3
    DONE = 4


class QuestionCellFormatter(CellFormatter):
    """Formats a cell containing a Science Bowl Question."""

    def __init__(
        self,
        force_capitalize: bool = False,
        line_after_stem: bool = False,
    ):
        self.force_capitalize = force_capitalize
        self.line_after_stem = line_after_stem

    def format(self, cell: _Cell) -> _Cell:
        """Takes a preprocessed question cell and returns a cell containing a
        properly-formatted Science Bowl question.

        Parameters
        ----------
        cell : _Cell

        Returns
        -------
        _Cell
            Cell containing a formatteed Science Bowl question.
        """

        state = QuestionFormatterState.Q_START
        q_type = None
        q_type_run = None
        current_choice = 0
        choices_para = {}

        for para in cell.paragraphs:
            if state is QuestionFormatterState.DONE:
                break

            elif state is QuestionFormatterState.Q_START and Q_TYPE_RE.match(para.text):
                try:
                    run_match = _validate_element_text(
                        q_type_run := para.runs[0], pattern=Q_TYPE_RE
                    )
                except QuestionParserException as ex:
                    highlight_cell_text(cell, WD_COLOR_INDEX.RED)
                    logger.error(str(ex))
                    break

                q_type = _format_question_type_run(q_type_run, run_match)

                if len(para.runs) == 1:
                    try:
                        _combine_qtype_and_stem_paragraphs(para)
                    except QuestionParserException as ex:
                        highlight_cell_text(cell, WD_COLOR_INDEX.RED)
                        logger.error(str(ex))
                        break

                # left pad the first run of the stem
                _left_pad_stem(stem_run=para.runs[1])

                state = QuestionFormatterState.STEM_END

            elif state is QuestionFormatterState.STEM_END:
                # handle incorrectly labeled questions and divert to the proper state
                if CHOICES_RE.match(para.text):
                    if q_type is QuestionType.SHORT_ANSWER:
                        logger.warning("Question type is SA, but has choices.")
                        q_type = _toggle_q_type_and_warn(q_type, q_type_run)
                    state = QuestionFormatterState.CHOICES

                elif ANSWER_RE.match(para.text):
                    if q_type is QuestionType.MULTIPLE_CHOICE:
                        logger.warning("Question type is MC, but has no choices.")
                        q_type = _toggle_q_type_and_warn(q_type, q_type_run)
                    state = QuestionFormatterState.ANSWER

            # this is intentionally an if - stem_end should continue onto
            # choices or answer
            if state is QuestionFormatterState.CHOICES and CHOICES_RE.match(para.text):
                try:
                    run_match = _validate_element_text(
                        (choice_run := para.runs[0]), pattern=CHOICES_RE
                    )
                except QuestionParserException as ex:
                    highlight_cell_text(cell, WD_COLOR_INDEX.RED)
                    logger.error(str(ex))
                    break

                _format_choice(run_match, choice_run, current_choice)

                # if current_choice was 0 and line_after_stem is true, we need to
                # insert a blank line before the first choice
                if current_choice == 0 and self.line_after_stem:
                    para.insert_paragraph_before("")

                # save text and update the choice we're looking for
                choices_para[current_choice] = para
                current_choice += 1

                if current_choice == 4:
                    state = QuestionFormatterState.ANSWER

            elif state is QuestionFormatterState.ANSWER and ANSWER_RE.match(para.text):
                para.insert_paragraph_before("")

                if self.force_capitalize:
                    capitalize_paragraph(para)

                if q_type is QuestionType.MULTIPLE_CHOICE:
                    try:
                        test_choice_match = _validate_element_text(para, TEST_CHOICE_RE)
                    except QuestionParserException as ex:
                        highlight_cell_text(cell, WD_COLOR_INDEX.RED)
                        logger.error(str(ex))
                        break

                    try:
                        _format_answer_line(para, test_choice_match, choices_para)
                    except QuestionParserException as ex:
                        logger.error(str(ex))

                state = QuestionFormatterState.DONE

        if state is not QuestionFormatterState.DONE:
            highlight_cell_text(cell, WD_COLOR_INDEX.RED)
            logger.error(f"Parsing failed while looking for {state}")

        return cell


class TuBCellFormatter(CellFormatter):
    def format(self, cell: _Cell) -> _Cell:
        if not (tub_match := TUB_RE.match(cell.text)):
            highlight_cell_text(cell, WD_COLOR_INDEX.RED)
            logger.error("Question must be a toss-up, bonus, or visual bonus.")

        else:
            put = TossUpBonus.from_string(tub_match.group(1)).value
            clear_cell(cell)
            tub_run = cell.paragraphs[0].runs[0]
            tub_run.text = put
            tub_run.italic = None
            tub_run.bold = None
            highlight_cell_text(cell, None)

        return cell


class SubjectCellFormatter(CellFormatter):
    def format(self, cell: _Cell) -> _Cell:
        subject_match = SUBJECT_RE.match(cell.text)

        if subject_match:
            put = Subject.from_string(subject_match.group(1)).value
            clear_cell(cell)
            subj_run = cell.paragraphs[0].runs[0]
            subj_run.text = put
            subj_run.italic = None
            subj_run.bold = None
            highlight_cell_text(cell, None)
        # if a match can't be found, highlight the cell red
        else:
            highlight_cell_text(cell, WD_COLOR_INDEX.RED)
            logger.error("Invalid subject.")

        return cell


class DifficultyFormatter(CellFormatter):
    """Formats the difficulty cell and shades it with the appropriate color."""

    color = COL_COLORS["LOD"]

    def format(self, cell: _Cell) -> _Cell:
        if cell.text:
            try:
                int(cell.text)
            except ValueError:
                highlight_cell_text(cell, WD_COLOR_INDEX.RED)
                logger.error("LOD should be blank or an integer.")

        return cell


class SetFormatter(CellFormatter):
    """Formats the set cell and shades it with the appropriate color."""

    color = COL_COLORS["Set"]


class RdFormatter(CellFormatter):
    """Formats the round cell and shades it with the appropriate color."""

    color = COL_COLORS["Rd"]


class QLetterFormatter(CellFormatter):
    """Formats the question letter cell and shades it with the appropriate color."""

    color = COL_COLORS["Q Letter"]


def _format_question_type_run(q_type_run: Run, run_match: re.Match) -> QuestionType:
    """Returns the type (Multiple Choice or Short Answer) of the question.

    Also handles italicizing the run containing the question type."""
    _q_type = QuestionType.from_string(run_match.group(1))
    # if the run contains more than the question type, split
    # the run into two
    if run_match.span()[1] < len(q_type_run.text):
        q_type_run, _ = split_run_at(q_type_run, run_match.span()[1])

    q_type_run.text, q_type_run.italic = _q_type.value, True
    return _q_type


def _format_answer_line(
    para: Paragraph, test_choice_match: re.Match, choices: Dict[int, Paragraph]
):
    choice_num = CHOICES.index(test_choice_match.group(1).upper() + ")")
    # if answer line is a single letter with an optional ), copy the
    # text of the correct choice over to the answer line
    if test_choice_match.span()[1] <= test_choice_match.span(1)[1] + 1:
        correct_para = choices[choice_num]
        para.text = "ANSWER: "
        for run in correct_para.runs:
            run_copy = deepcopy(run._r)
            run_copy.text = run_copy.text.upper()
            para._p.append(run_copy)
        fuse_consecutive_runs(para)
    # otherwise, check that the answer line text matches the choice.
    # if it doesn't raise a linting error.
    else:
        if (
            choices[choice_num].text.upper()
            != para.text.replace("ANSWER: ", "", 1).upper()
        ):
            highlight_paragraph_text(para, WD_COLOR_INDEX.YELLOW)
            raise QuestionParserException("Answer line doesn't match choice.")


def _validate_element_text(
    run_or_para: Union[Run, Paragraph], pattern: re.Pattern
) -> re.Match:
    """If a recognized part of a question (question type, choices, etc.) is split across
    multiple runs, it won't parse."""

    msgs = {
        TEST_CHOICE_RE: "Found answer line, but couldn't find W, X, Y, or Z.",
        CHOICES_RE: "Couldn't parse question. Check that choices are correct.",
        Q_TYPE_RE: "Couldn't parse question. Check that question type is correct.",
    }

    if not (run_match := pattern.match(run_or_para.text)):
        raise QuestionParserException(msgs[pattern])

    return run_match


def _combine_qtype_and_stem_paragraphs(para: Paragraph) -> None:
    """If a paragraph has only one run, it should be combined with the next paragraph,
    which is presumably the stem paragraph.

    Raises an error of W) or ANSWER: are found in the next paragraph, however."""
    next_para = para._p.getnext()
    next_para_text = "".join(_r.text for _r in next_para if _r.text)
    # need to check if there even IS a stem - next paragraph shouldn't
    # start with W) or ANSWER:
    if CHOICES_RE.match(next_para_text) or ANSWER_RE.match(next_para_text):
        raise QuestionParserException("Couldn't parse question.")

    move_runs_to_end_of_para(next_para, para._p)


def _left_pad_stem(stem_run: Run):
    """All question stems should start with four spaces."""
    stem_run.text = f"    {stem_run.text.lstrip()}"


def _toggle_q_type_and_warn(q_type: QuestionType, q_type_run: Run) -> QuestionType:
    """If a question type is mislabeled, highlight it yellow and toggle the question
    type."""
    q_type_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if q_type is QuestionType.SHORT_ANSWER:
        return QuestionType.MULTIPLE_CHOICE
    elif q_type is QuestionType.MULTIPLE_CHOICE:
        return QuestionType.SHORT_ANSWER


def _format_choice(run_match: re.Match, choice_run: Run, current_choice: int):
    """If the wrong choice was matched, replace it with the right choice."""
    if run_match.group(1) != CHOICES[current_choice]:
        choice_run.text = choice_run.text.replace(
            run_match.group(1), CHOICES[current_choice], 1
        )
