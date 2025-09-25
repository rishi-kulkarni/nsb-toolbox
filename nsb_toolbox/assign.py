import re
from functools import reduce
from typing import Optional
from typing_extensions import Self

import docx.table
import docx.document
import numpy as np
from scipy.optimize import linear_sum_assignment


from ._base_questions import BaseScienceBowlQuestions
from .classes import TossUpBonus
from .tables import COL_MAPPING
from .yamlparsers import ParsedQuestionSpec


class EditedQuestions(BaseScienceBowlQuestions):
    """Represents an edited set of Science Bowl questions from a single subject
    with level of difficulty and subcategory assigned by the SME.

    Parameters
    ----------
    document : docx.document.Document
        Word document of Science Bowl questions wrapped by this class.

    Raises
    ------
    ValueError
        Raised if there are formatting issues in the document.

    Attributes
    ----------
    tubs : np.ndarray[str]
    difficulties : np.ndarray[int]
    qtypes : np.ndarray[str]
    subcategories : np.ndarray[str]
    sets : List[docx.table._Cell]
    rounds : List[docx.table._Cell]
    qletters : List[docx.table._Cell]

    Methods
    -------
    assign(question_spec: ParsedQuestionSpec)
        Assigns the questions to a specification via a linear sum assignment.
    """

    def __init__(self, document: docx.document.Document) -> Self:

        super().__init__(document)

        try:
            self.difficulties
            self.tubs
            self.qtypes
        except ValueError as ex:
            raise ValueError(
                "One or more issues with the question document."
                " Have you run nsb format on it?\n",
                ex,
            )

    def assign(
        self, question_spec: ParsedQuestionSpec, dry_run: Optional[bool] = False
    ):
        """Attempts to find a valid assignment of the questions to question_spec via a
        linear sum assignment. Upon success, writes the successful assignment to the
        question document. Otherwise, notifies the user which parts of the specification
        could not be met.

        Parameters
        ----------
        question_spec : ParsedQuestionSpec
            Question specification read from a config file.
        dry_run : Optional[bool]
            If True, will not write assignments back to question file, by default False.

        Raises
        ------
        ValueError
            Raised if the questions cannot meet the given specification. Prints
            the specifications that cannot be met so that the SME can find appropriate
            questions.
        """
        self._validate(question_spec)
        self._report_stats(question_spec)

        cost_matrix = build_cost_matrix(questions=self, spec=question_spec)

        q_assignments, round_assignments = linear_sum_assignment(cost_matrix)
        assignment_costs = cost_matrix[q_assignments, round_assignments]

        if assignment_costs.sum() > 1_000_000:
            self._raise_assignment_failure(question_spec, assignment_costs)

        else:
            self._write_assignment(
                question_spec, q_assignments, round_assignments, dry_run
            )

    def _write_assignment(
        self,
        question_spec: ParsedQuestionSpec,
        q_assignments: np.ndarray,
        round_assignments: np.ndarray,
        dry_run: Optional[bool] = False,
    ):
        """In case of assignment success, write successful assignemnt to
        this question set."""
        if not dry_run:
            print("\nFound a successful set of assignments!")
            for q_idx, spec_idx in zip(q_assignments, round_assignments):
                self.sets[q_idx].text = question_spec.sets[spec_idx]
                self.rounds[q_idx].text = question_spec.rounds[spec_idx]
                self.qletters[q_idx].text = question_spec.qletters[spec_idx]
        else:
            print("\nNot writing assignments as this is a dry run.")

    def sort_assignments(self, question_spec: ParsedQuestionSpec) -> None:
        """Reorders the question rows to follow the YAML configuration order."""
        if not self.document.tables:
            return

        table = self.document.tables[0]
        data_rows = list(table.rows)[1:]

        if not data_rows:
            return

        tub_lookup = {
            TossUpBonus.TOSS_UP.value: 0,
            TossUpBonus.BONUS.value: 1,
            TossUpBonus.VISUAL_BONUS.value: 2,
        }

        assigned_rows = []
        unassigned_rows = []

        for idx, row in enumerate(data_rows):
            tub_text = row.cells[COL_MAPPING["TUB"]].text.strip()
            set_text = row.cells[COL_MAPPING["Set"]].text.strip()
            round_text = row.cells[COL_MAPPING["Rd"]].text.strip()
            letter_text = row.cells[COL_MAPPING["Q Letter"]].text.strip()

            if not (tub_text and set_text and round_text and letter_text):
                unassigned_rows.append((idx, row))
                continue

            try:
                tub_value = TossUpBonus.from_string(tub_text).value
            except ValueError:
                unassigned_rows.append((idx, row))
                continue

            assigned_rows.append(
                {
                    "row": row,
                    "set": set_text,
                    "round": round_text,
                    "letter": letter_text.upper(),
                    "tub": tub_value,
                    "original_index": idx,
                }
            )

        if not assigned_rows:
            return

        set_order = {}
        prefix_order = {}
        round_order = {}
        letter_order = {}

        round_pattern = re.compile(r"^([^\d]*)(\d+)(.*)$")

        def decompose_round(value: str):
            match = round_pattern.match(value)
            if match:
                prefix, number, suffix = match.groups()
                return prefix, int(number), suffix
            return value, float("inf"), ""

        for detail in question_spec.question_list:
            set_order.setdefault(detail.set, len(set_order))

            prefix, number, suffix = decompose_round(detail.round)
            prefix_map = prefix_order.setdefault(detail.set, {})
            prefix_idx = prefix_map.setdefault(prefix, len(prefix_map))

            round_map = round_order.setdefault(detail.set, {})
            round_map.setdefault(
                detail.round,
                (prefix_idx, number, suffix, len(round_map)),
            )

            letter_order.setdefault(detail.letter.upper(), len(letter_order))

        def round_sort(info):
            set_rounds = round_order.get(info["set"], {})
            if info["round"] in set_rounds:
                return set_rounds[info["round"]]

            prefix, number, suffix = decompose_round(info["round"])
            prefix_idx = prefix_order.get(info["set"], {}).get(
                prefix,
                len(prefix_order.get(info["set"], {})),
            )
            return (prefix_idx, number, suffix, info["original_index"])

        def sort_key(info):
            set_idx = set_order.get(info["set"], len(set_order))
            round_idx = round_sort(info)
            letter_idx = letter_order.get(
                info["letter"], len(letter_order) + info["original_index"]
            )
            tub_idx = tub_lookup.get(info["tub"], len(tub_lookup))
            return (set_idx, round_idx, letter_idx, tub_idx, info["original_index"])

        assigned_rows.sort(key=sort_key)

        tbl = table._tbl
        for row in data_rows:
            tbl.remove(row._tr)

        for info in assigned_rows:
            tbl.append(info["row"]._tr)

        for _, row in sorted(unassigned_rows, key=lambda item: item[0]):
            tbl.append(row._tr)

        self._cells = self.document.tables[0]._cells
        for attr in (
            "tubs",
            "difficulties",
            "qtypes",
            "subcategories",
            "writers",
            "stats",
        ):
            self.__dict__.pop(attr, None)

    def _raise_assignment_failure(
        self, question_spec: ParsedQuestionSpec, assignment_costs: np.ndarray
    ):
        """In case assignment fails, find the questions in the spec that failed to
        be assigned and return them to the user."""
        where_bad_assigns = np.argwhere(
            assignment_costs == assignment_costs.max()
        ).ravel()
        failed_assignments = [
            question_spec.question_list[idx] for idx in where_bad_assigns
        ]

        failed_sets = (question.set for question in failed_assignments)
        failed_diff = (question.difficulty for question in failed_assignments)
        failed_tubs = (question.tub for question in failed_assignments)
        failed_qtypes = (question.qtype for question in failed_assignments)

        raw_values = np.array(
            [
                f"{set_:<10}{difficulty:^5}{tub.value:^5}{qtype.value:>20}"
                if qtype is not None
                else f"{set_:<10}{difficulty:^5}{tub.value:^5}{'Any':>20}"
                for set_, difficulty, tub, qtype in zip(
                    failed_sets, failed_diff, failed_tubs, failed_qtypes
                )
            ]
        )

        failed_stats = {
            val: count for val, count in zip(*np.unique(raw_values, return_counts=True))
        }

        failed_table = "\n".join(
            f"{key}{failed_stats.get(key, 0):>10}" for key in failed_stats.keys()
        )
        raise ValueError(
            "Failed to assign. Do you have enough questions?\n"
            "Missing questions:\n"
            f"{'-':->60}\n"
            f"{'Set':<10}{'LOD':^4}{'TUB':^6}{'QType':^26}{'Need':>9}\n"
            f"{'-':->60}\n"
            f"{failed_table}"
        )

    def _report_stats(self, question_spec: ParsedQuestionSpec):

        print("\nStatistics")
        print(f"{'-':->43}")
        print(f"{'Set':<9}{'LOD':^}{'TUB':>10}{'Need':>11}{'Have':>10}")
        print(f"{'-':->43}")

        for key in sorted(self.stats.keys() | question_spec.stats.keys()):

            need = question_spec.stats.get(key, 0)
            have = self.stats.get(key, 0)

            print(f"{key}{need:>10}{have:>10}")

    def _validate(self, question_spec: ParsedQuestionSpec):
        """Rounds and question letters in self.document should be empty. If not,
        ask the user to confirm that they intend to overwrite them.

        Also, there should be more questions available than are required by the spec."""

        if len(self.difficulties) < len(question_spec.question_list):
            raise ValueError(
                "There are not enough available questions"
                " in this document to fill the specified rounds."
            )

        if any(x.text for x in self.rounds) or any(x.text for x in self.qletters):
            while True:
                user_input = input(
                    "The Round and Q Letter columns are not empty in this document.\n"
                    "Continuing will overwrite the current assignments. Continue? (Y/n)"
                )
                if user_input.lower() in ("y", "n"):
                    break

            if user_input.lower() == "n":
                raise ValueError("Aborted!")
            else:
                for round_ in self.rounds:
                    round_.text = ""
                for qletter in self.qletters:
                    qletter.text = ""


def build_cost_matrix(
    questions: EditedQuestions,
    spec: ParsedQuestionSpec,
) -> np.ndarray:
    """Computes a cost matrix for assigning a set of Science Bowl questions to a
    given question specification.

    Parameters
    ----------
    questions : EditedQuestions
    spec : ParsedQuestionSpec


    Returns
    -------
    np.ndarray
        Cost matrix with questions represented by rows and slots by columns.
    """
    # randomness, can be seeded in the config file
    random_matrix = spec.config.rng.uniform(
        0,
        0.0001,
        size=(questions.difficulties.size, spec.difficulties.size),
    )

    # squared loss for difficulties
    diff_matrix = (spec.difficulties - questions.difficulties[:, np.newaxis]) ** 2

    # penalize subcategory mismatches
    subcat_matrix = np.where(
        (spec.subcategories != questions.subcategories[:, np.newaxis])
        & (spec.subcategories != ""),
        spec.config.subcat_mismatch_penalty,
        0,
    )

    # penalize unpreferred writers
    if spec.config.preferred_writers:
        writer_matrix = np.where(
            (
                questions.writers[:, np.newaxis]
                != np.array(spec.config.preferred_writers)
            ).all(axis=1),
            0.1,
            0,
        )[:, np.newaxis]
    else:
        writer_matrix = np.zeros_like(diff_matrix)

    cost_matrix = random_matrix + diff_matrix + subcat_matrix + writer_matrix
    invalid_assignments = invalid_assignment_mask(questions, spec)
    # invalid assignments need a finite but large cost
    cost_matrix[invalid_assignments] = 1_000_000

    return cost_matrix


def invalid_assignment_mask(
    questions: EditedQuestions, spec: ParsedQuestionSpec
) -> np.ndarray:
    """Applies the following rules to ensure that no invalid assignments are made:

    1. Questions with missing values in the LOD column cannot be assigned.
    2. TOSS-UP or BONUS markings should be respected.
    3. If a question must be Short Answer, it should not have a Multiple Choice
        question assigned to it.
    4. If a question is marked with the "B" set, it should not be assigned to the
        "A" set. However, questions that lack an "-A" or "-B" masking can be
        assigned to either set.

    Parameters
    ----------
    questions : EditedQuestions
    spec : ParsedQuestionSpec

    Returns
    -------
    np.ndarray
        2D Boolean mask for the cost matrix where True indicates an invalid assignment.
    """
    masks = []
    # mask to indicate where question LODs are missing
    masks.append(questions.difficulties[:, np.newaxis] == -1)

    # mask to indicate where toss-up/bonus do not match
    masks.append(spec.tubs != questions.tubs[:, np.newaxis])

    # mask to indicate where Short Answer/Multiple Choice do not match
    masks.append((spec.qtypes != questions.qtypes[:, np.newaxis]) & (spec.qtypes != ""))

    # mask to indicate where Sets match, if the question has a Set indicated
    q_sets = np.array([x.text for x in questions.sets])[:, np.newaxis]
    packets = np.unique([x.split("-")[0] for x in spec.sets])
    masks.append((spec.sets != q_sets) & (q_sets != packets).all(axis=1)[:, np.newaxis])

    # apply all masks
    return reduce(np.logical_or, masks)
